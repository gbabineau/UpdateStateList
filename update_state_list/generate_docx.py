"""
Main function for the generate_docx application
plumages.
"""

import argparse
import csv
import logging
import tomllib

from docx import Document, opc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches


def add_hyperlink(paragraph, url, text, color, underline):
    """
    Add a hyperlink to a paragraph in a Word document.

    Args:
        paragraph: The paragraph object to add the hyperlink to.
        url (str): The URL that the hyperlink should point to.
        text (str): The display text for the hyperlink.
        color (str, optional): The color of the hyperlink text
        underline (bool): Whether the hyperlink should be underlined.

    Returns:
        OxmlElement: The hyperlink XML element that was created and added to
        the paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rpr.append(c)

    if not underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "none")
        rpr.append(u)

    new_run.append(rpr)
    new_run.text = text
    hyperlink.append(new_run)
    # pylint: disable=W0212
    paragraph._p.append(hyperlink)

    return hyperlink


def generate_docx(official_list_file) -> None:
    """
    Generate a formatted Word document from a CSV file containing official bird
    species data.

    Args:
        official_list_file: Path to the CSV file containing bird species data.
    """
    # Read CSV data
    with open(official_list_file, "r", encoding="utf-8") as f:
        birds_data = list(csv.DictReader(f))

    # Initialize document
    doc = Document()
    title = doc.add_heading(
        "The Birds of Virginia and its Offshore Waters: The Official List", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"

    # Set headers
    headers = [
        "#",
        "Species",
        "Scientific Name",
        "State Status",
        "Spatial Distribution",
        "Counts & Seasonality",
    ]
    widths = [0.5, 1.1, 1.1, 1.1, 1.1, 1.1]

    for i, (header, width) in enumerate(zip(headers, widths)):
        table.columns[i].width = Inches(width)
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Add data rows
    current_order = current_family = ""
    index = 1
    historically_occurring_section = False
    for bird in birds_data:
        # Add historical species row if first occurrence
        state_status = bird.get("State Status", "")
        if state_status == "(4)" and not historically_occurring_section:
            _add_category_row(
                table,
                "Species Believed to Have Occurred Historically",
                "FFFFE0",
            )
            historically_occurring_section = True
        # Add order row if changed
        if bird.get("order", "") != current_order:
            current_order = bird["order"]
            current_family = ""
            _add_category_row(table, f"Order: {current_order}", "D3D3D3")

        # Add family row if changed
        if bird.get("familyComName", "") != current_family:
            current_family = bird["familyComName"]
            _add_category_row(table, f"Family: {current_family}", "ADD8E6")

        # Add species row
        row_cells = table.add_row().cells
        species_code = bird.get("speciesCode", "")

        if (
            bird.get("subspecies", "False").lower() == "false"
            and not historically_occurring_section
        ):
            row_cells[0].text = str(index)
            index += 1

        add_hyperlink(
            row_cells[1].paragraphs[0],
            f"https://ebird.org/species/{species_code}/US-VA",
            bird.get("comName"),
            "0000FF",
            False,
        )

        row_cells[2].text = bird.get("sciName", "")
        row_cells[3].text = state_status

        add_hyperlink(
            row_cells[4].paragraphs[0],
            f"http://ebird.org/ebird/map/{species_code}?neg=true&env.minX="
            "-84.70&env.minY=36.20&env.maxX=-70.95&env.maxY=37.22&zh=true&"
            "gp=true&ev=Z&mr=1-12&bmo=1&emo=12&yr=all&getLocations=states&"
            "states=US-VA",
            "Map",
            "0000FF",
            False,
        )

        add_hyperlink(
            row_cells[5].paragraphs[0],
            f"http://ebird.org/ebird/GuideMe?cmd=decisionPage&speciesCodes="
            f"{species_code}&getLocations=states&states=US-VA&bYear=1900&eYear="
            "Cur&bMonth=1&eMonth=12&reportType=species&parentState=US-VA",
            "Chart",
            "0000FF",
            False,
        )

    # Save document
    output_file = official_list_file.replace(".csv", ".docx")
    doc.save(output_file)
    logging.info("Document saved as %s", output_file)


def _add_category_row(table, text, color):
    """Add a category row (order/family) to the table."""
    row_cells = table.add_row().cells
    merged_cell = row_cells[0].merge(row_cells[5])
    merged_cell.text = text
    merged_cell.paragraphs[0].runs[0].bold = True
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    # pylint: disable=W0212
    merged_cell._element.get_or_add_tcPr().append(shading_elm)


def main():
    """Main function for the app."""
    arg_parser = argparse.ArgumentParser(
        prog="update-state-list", description="Update elements of a state list."
    )
    with open("pyproject.toml", "rb") as f:
        pyproject_data = tomllib.load(f)
    version = (
        pyproject_data.get("tool", {}).get("poetry", {}).get("version", "0.0.0")
    )
    arg_parser.add_argument(
        "--version", action="version", version=f"%(prog)s {version}"
    )
    arg_parser.add_argument(
        "--verbose", action="store_true", help="increase verbosity"
    )
    arg_parser.add_argument(
        "--official_list_csv",
        required=True,
        help="csv of official list created by update_state_list",
    )
    args = arg_parser.parse_args()

    if args.verbose:
        logging.basicConfig(level=logging.INFO)

    generate_docx(args.official_list_csv)


if __name__ == "__main__":
    main()
