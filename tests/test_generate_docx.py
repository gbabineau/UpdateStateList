"""Tests for generate_docx module."""

from unittest.mock import MagicMock, mock_open, patch

import pytest
from docx import Document
from docx.oxml import CT_Hyperlink

from update_state_list.generate_docx import (add_hyperlink,
                                             generate_docx,
                                             main,
                                             _add_category_row)


class TestAddHyperlink:
    """Tests for add_hyperlink function."""

    def test_add_hyperlink_with_color_and_underline(self):
        """Test adding hyperlink with color and underline."""
        doc = Document()
        paragraph = doc.add_paragraph()

        hyperlink = add_hyperlink(
            paragraph,
            "https://example.com",
            "Test Link",
            "FF0000",
            True
        )

        assert hyperlink is not None
        assert isinstance(hyperlink, CT_Hyperlink)

    def test_add_hyperlink_without_underline(self):
        """Test adding hyperlink without underline."""
        doc = Document()
        paragraph = doc.add_paragraph()

        hyperlink = add_hyperlink(
            paragraph,
            "https://example.com",
            "Test Link",
            "0000FF",
            False
        )

        assert hyperlink is not None

    def test_add_hyperlink_without_color(self):
        """Test adding hyperlink without color."""
        doc = Document()
        paragraph = doc.add_paragraph()

        hyperlink = add_hyperlink(
            paragraph,
            "https://example.com",
            "Test Link",
            None,
            False
        )

        assert hyperlink is not None


class TestGenerateDocx:
    """Tests for generate_docx function."""

    @pytest.fixture
    def sample_csv_data(self):
        """Sample CSV data for testing."""
        return [
            {
                "speciesCode": "gockin",
                "order": "Passeriformes",
                "familyComName": "Kinglets",
                "comName": "Golden-crowned Kinglet",
                "sciName": "Regulus satrapa",
                "State Status": "Common",
                "subspecies": "False"
            },
            {
                "speciesCode": "rucroc",
                "order": "Passeriformes",
                "familyComName": "Kinglets",
                "comName": "Ruby-crowned Kinglet",
                "sciName": "Corthylio calendula",
                "State Status": "Common",
                "subspecies": "False"
            }
        ]

    @patch('update_state_list.generate_docx.Document')
    @patch("update_state_list.generate_docx.add_hyperlink")
    @patch('builtins.open', new_callable=mock_open)
    @patch('csv.DictReader')
    def test_generate_docx_creates_document(self, mock_csv, _, __ , mock_doc):
        """Test that generate_docx creates a document."""
        mock_csv.return_value = [
            {
                "speciesCode": "gockin",
                "order": "Passeriformes",
                "familyComName": "Kinglets",
                "comName": "Golden-crowned Kinglet",
                "sciName": "Regulus satrapa",
                "State Status": "Common",
                "subspecies": "False"
            }
        ]

        mock_doc_instance = MagicMock()
        mock_doc.return_value = mock_doc_instance

        generate_docx("test.csv")

        mock_doc_instance.save.assert_called_once_with("test.docx")

    @patch('update_state_list.generate_docx.Document')
    @patch("update_state_list.generate_docx.add_hyperlink")
    @patch('builtins.open', new_callable=mock_open)
    @patch('csv.DictReader')
    def test_generate_docx_handles_subspecies(self, mock_csv, _, __, mock_doc):
        """Test that subspecies are handled correctly."""
        mock_csv.return_value = [
            {
                "speciesCode": "gockin",
                "order": "Passeriformes",
                "familyComName": "Kinglets",
                "comName": "Golden-crowned Kinglet",
                "sciName": "Regulus satrapa",
                "State Status": "Common",
                "subspecies": "True"
            }
        ]

        mock_doc_instance = MagicMock()
        mock_doc.return_value = mock_doc_instance

        generate_docx("test.csv")

        mock_doc_instance.save.assert_called_once()


class TestMain:
    """Tests for main function."""

    @patch('update_state_list.generate_docx.generate_docx')
    @patch('builtins.open', new_callable=mock_open, read_data=b'[tool.poetry]\nversion = "1.0.0"\n')
    @patch('sys.argv', ['prog', '--official_list_csv', 'test.csv'])
    def test_main_with_required_args(self, _, mock_generate):
        """Test main function with required arguments."""
        main()
        mock_generate.assert_called_once_with('test.csv')

class TestAddCategoryRow:
    """Tests for _add_category_row function."""

    def test_add_category_row_creates_merged_cell(self):
        """Test that _add_category_row creates a merged cell."""

        doc = Document()
        table = doc.add_table(rows=1, cols=6)

        _add_category_row(table, "Order: Passeriformes", "D3D3D3")

        # Should have 2 rows now (header + category row)
        assert len(table.rows) == 2

    def test_add_category_row_sets_text(self):
        """Test that _add_category_row sets the correct text."""

        doc = Document()
        table = doc.add_table(rows=1, cols=6)

        _add_category_row(table, "Family: Kinglets", "ADD8E6")

        merged_cell = table.rows[1].cells[0]
        assert merged_cell.text == "Family: Kinglets"

    def test_add_category_row_makes_text_bold(self):
        """Test that _add_category_row makes text bold."""

        doc = Document()
        table = doc.add_table(rows=1, cols=6)

        _add_category_row(table, "Order: Test", "D3D3D3")

        merged_cell = table.rows[1].cells[0]
        assert merged_cell.paragraphs[0].runs[0].bold is True

    def test_add_category_row_with_different_colors(self):
        """Test that _add_category_row works with different colors."""

        doc = Document()
        table = doc.add_table(rows=1, cols=6)

        # Test with order color
        _add_category_row(table, "Order: Test", "D3D3D3")
        assert len(table.rows) == 2

        # Test with family color
        _add_category_row(table, "Family: Test", "ADD8E6")
        assert len(table.rows) == 3
